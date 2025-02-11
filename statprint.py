from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import pandas as pd

class StatPrint:
    def __init__(self, filename="report", doc_type="word", title="Report"):
        # Initialize with default filename, doc_type, and title
        self.filename = filename
        self.doc_type = doc_type
        self.title = title  # Set the title
        self.content = []
        self.graph_count = 0  # To create unique filenames for graphs

    def add_heading(self, heading):
        """Add a heading to the report."""
        self.content.append(('heading', heading))

    def add_table(self, data, custom_headers=None, indent_rows=None):
        """
        Add a table to the report.
        
        Parameters:
          - data: a pandas DataFrame or Series.
          - custom_headers: Optional list of column names to override the default.
          - indent_rows: Optional list of row indices (starting at 0 for the first data row)
                         that should have their first cell indented.
        """
        if indent_rows is None:
            indent_rows = []

        if isinstance(data, pd.Series):
            # For Series (often from value_counts), reset the index so you have two columns.
            df = data.reset_index()
            if custom_headers is None:
                if data.name is not None:
                    custom_headers = [data.name, "Count"]
                else:
                    custom_headers = ["Value", "Count"]
            df.columns = custom_headers
        else:
            df = data.copy()
            if custom_headers is not None:
                df.columns = custom_headers

        # Store the table as a dictionary so that additional styling options can be used later.
        self.content.append(('table', {'df': df, 'custom_headers': custom_headers, 'indent_rows': indent_rows}))

    def add_graph(self, graph, filename=None):
        """
        Save the graph (a matplotlib figure) to a unique file and add its filename to the report.
        
        Parameters:
          - graph: a matplotlib figure object.
          - filename: Optional filename base; if provided the counter is prepended.
        """
        if filename is None:
            filename = f"graph_{self.graph_count}.png"
        else:
            filename = f"{self.graph_count}_{filename}"
        self.graph_count += 1
        graph.savefig(filename, format='png')
        self.content.append(('graph', filename))

    def generate_report(self):
        if self.doc_type == 'word':
            self.generate_word_report()
        elif self.doc_type == 'pdf':
            self.generate_pdf_report()

    def _apply_word_table_style(self, table, indent_rows):
        """
        Helper method to modify a docx table:
          - Makes header text bold.
          - Removes vertical borders (leaving horizontal borders).
          - Applies left indentation to rows specified in indent_rows (data rows only).
        """
        # Bold header row (assumed to be the first row)
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Remove vertical borders but keep horizontal borders.
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn

        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblBorders = parse_xml(
            r'<w:tblBorders %s>'
            r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:left w:val="nil"/>'
            r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:right w:val="nil"/>'
            r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:insideV w:val="nil"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr.append(tblBorders)

        # Indent specified data rows.
        # Note: table.rows[0] is the header, so data rows start at index 1.
        for i, row in enumerate(table.rows[1:]):
            if i in indent_rows:
                cell = row.cells[0]
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.left_indent = Inches(0.25)

    def generate_word_report(self):
        doc = Document()
        doc.add_heading(self.title, 0)

        for content_type, content in self.content:
            if content_type == 'heading':
                doc.add_heading(content, level=1)
            elif content_type == 'table':
                table_info = content
                df = table_info['df']
                indent_rows = table_info.get('indent_rows', [])
                # Create a table with one header row and as many columns as df has
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'

                # Add headers
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(df.columns):
                    hdr_cells[i].text = str(col_name)
                    # Make header text bold (in case style doesn't do it)
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Add data rows.
                for idx, row in enumerate(df.itertuples(index=False, name=None)):
                    row_cells = table.add_row().cells
                    for j, value in enumerate(row):
                        row_cells[j].text = str(value)
                        # If this row is designated to be indented, indent the first cell.
                        if j == 0 and idx in indent_rows:
                            for paragraph in row_cells[j].paragraphs:
                                paragraph.paragraph_format.left_indent = Inches(0.25)
                # Apply the border style modifications.
                self._apply_word_table_style(table, indent_rows)
            elif content_type == 'graph':
                doc.add_paragraph()  # add spacing before image
                doc.add_picture(content, width=Inches(6))  # adjust width as needed

        doc.save(f"{self.filename}.docx")
        print(f"Report saved as {self.filename}.docx")

    def generate_pdf_report(self):
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, txt=self.title, ln=True, align='C')
        pdf.ln(5)

        pdf.set_font("Arial", '', 12)
        for content_type, content in self.content:
            if content_type == 'heading':
                pdf.ln(10)
                pdf.set_font("Arial", 'B', 14)
                pdf.cell(0, 10, txt=content, ln=True)
                pdf.set_font("Arial", '', 12)
            elif content_type == 'table':
                table_info = content
                df = table_info['df']
                indent_rows = table_info.get('indent_rows', [])

                # Calculate effective page width.
                effective_width = pdf.w - 2 * pdf.l_margin
                col_width = effective_width / len(df.columns)
                
                # Header row (bold, no cell borders; draw a horizontal line below)
                pdf.set_font("Arial", 'B', 10)
                for col in df.columns:
                    pdf.cell(col_width, 10, str(col), border=0, align='C')
                pdf.ln(10)
                # Draw horizontal line below header.
                pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + effective_width, pdf.get_y())

                # Data rows.
                pdf.set_font("Arial", '', 10)
                for i, row in enumerate(df.itertuples(index=False, name=None)):
                    for j, value in enumerate(row):
                        text = str(value)
                        # Indent the first column if this row is marked as a subclass.
                        if j == 0 and i in indent_rows:
                            text = "    " + text  # adjust the number of spaces as needed
                        pdf.cell(col_width, 10, text, border=0)
                    pdf.ln(10)
                    # Draw horizontal line after each row.
                    pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + effective_width, pdf.get_y())
            elif content_type == 'graph':
                pdf.ln(10)
                pdf.image(content, x=None, y=None, w=150)  # Adjust width as necessary

        pdf.output(f"{self.filename}.pdf")
        print(f"Report saved as {self.filename}.pdf")
